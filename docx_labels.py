##############################################################################
# docx_labels.py
#
# Avery 94207 Word Label Generator
# Version 3
#
# Cochise County Master Gardener Association
##############################################################################

from io import BytesIO

from docx import Document

from docx.enum.section import WD_SECTION
from docx.enum.table import (
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
    WD_CELL_VERTICAL_ALIGNMENT,
)
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
)
from docx.shared import (
    Inches,
    Pt,
    RGBColor,
)

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


##############################################################################
# Colors
##############################################################################

GREEN = RGBColor(0, 102, 51)
RED   = RGBColor(180, 0, 0)
BLACK = RGBColor(0, 0, 0)


##############################################################################
# Font Sizes
##############################################################################

TITLE_SIZE   = 10
FAMILY_SIZE  = 11
VARIETY_SIZE = 10

BODY_SIZE    = 8
SMALL_SIZE   = 7
TINY_SIZE    = 6


##############################################################################
# Avery 94207 Geometry
##############################################################################

PAGE_WIDTH  = 8.50
PAGE_HEIGHT = 11.00

TOP_MARGIN    = 0.49
BOTTOM_MARGIN = 0.49
LEFT_MARGIN   = 0.125
RIGHT_MARGIN  = 0.25

LABEL_WIDTH   = 4.00
LABEL_HEIGHT  = 2.00

GUTTER_WIDTH  = 0.125
FILLER_WIDTH  = 0.125

ROWS = 5
COLS = 3

LABELS_PER_PAGE = 10


##############################################################################
# XML Helpers
##############################################################################

def _remove_table_borders(table):
    """
    Remove every visible border from a table.
    """

    tblPr = table._tbl.tblPr

    borders = tblPr.first_child_found_in("w:tblBorders")

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    for edge in (
        "top",
        "left",
        "bottom",
        "right",
        "insideH",
        "insideV",
    ):

        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)


def _set_row_height(row):
    """
    Force an exact Avery row height.
    """

    row.height = Inches(LABEL_HEIGHT)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def _set_cell_width(cell, width):
    """
    Force an exact cell width.
    """

    cell.width = Inches(width)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcW = tcPr.first_child_found_in("w:tcW")

    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)

    tcW.set(qn("w:type"), "dxa")
    tcW.set(
        qn("w:w"),
        str(int(width * 1440)),
    )


def _set_cell_margins(
    cell,
    top=0,
    bottom=0,
    left=0,
    right=0,
):
    """
    Remove Word's default cell padding.

    Values are twips.
    """

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = tcPr.first_child_found_in("w:tcMar")

    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for name, value in (
        ("top", top),
        ("left", left),
        ("bottom", bottom),
        ("right", right),
    ):

        node = tcMar.find(qn(f"w:{name}"))

        if node is None:
            node = OxmlElement(f"w:{name}")
            tcMar.append(node)

        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _clear_cell(cell):
    """
    Remove Word's default empty paragraph.
    """

    if cell.paragraphs:

        p = cell.paragraphs[0]._element
        p.getparent().remove(p)

##############################################################################
# Document Setup
##############################################################################

def _setup_document():
    """
    Create a Word document configured specifically for
    Avery 94207 labels.
    """

    document = Document()

    ######################################################################
    # Page setup
    ######################################################################

    section = document.sections[0]

    section.start_type = WD_SECTION.NEW_PAGE

    section.page_width  = Inches(PAGE_WIDTH)
    section.page_height = Inches(PAGE_HEIGHT)

    section.left_margin   = Inches(LEFT_MARGIN)
    section.right_margin  = Inches(RIGHT_MARGIN)
    section.top_margin    = Inches(TOP_MARGIN)
    section.bottom_margin = Inches(BOTTOM_MARGIN)

    ######################################################################
    # Header / Footer
    ######################################################################

    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    ######################################################################
    # Normal style
    ######################################################################

    normal = document.styles["Normal"]

    normal.font.name = "Arial"
    normal.font.size = Pt(BODY_SIZE)

    pf = normal.paragraph_format

    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)

    pf.line_spacing = 1.0

    pf.keep_together = True
    pf.keep_with_next = False

    ######################################################################
    # Remove the automatic empty paragraph
    ######################################################################

    if document.paragraphs:

        p = document.paragraphs[0]._element
        p.getparent().remove(p)

    return document


##############################################################################
# Paragraph / Text Helpers
##############################################################################

def _paragraph(
    container,
    align=WD_ALIGN_PARAGRAPH.LEFT,
):
    """
    Create a paragraph with ZERO extra spacing.
    """

    p = container.add_paragraph()

    fmt = p.paragraph_format

    fmt.space_before = Pt(0)
    fmt.space_after  = Pt(0)

    fmt.line_spacing = 1.0

    fmt.keep_together = True
    fmt.keep_with_next = False

    p.alignment = align

    return p


def _add_text(
    container,
    text,
    *,
    size=BODY_SIZE,
    bold=False,
    italic=False,
    color=BLACK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
):
    """
    Add one formatted paragraph.
    """

    if text is None:
        return

    text = str(text).strip()

    if not text:
        return

    p = _paragraph(
        container,
        align=align,
    )

    run = p.add_run(text)

    run.font.name = "Arial"
    run.font.size = Pt(size)

    run.bold = bold
    run.italic = italic

    run.font.color.rgb = color


def _add_field(
    container,
    label,
    value,
):
    """
    Add a compact right-side field such as:

        Year: 2026

    with the label in bold and the value in regular weight.
    """

    if value is None:
        return

    value = str(value).strip()

    if not value:
        return

    p = _paragraph(container)

    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(BODY_SIZE)

    r2 = p.add_run(value)
    r2.font.name = "Arial"
    r2.font.size = Pt(BODY_SIZE)


##############################################################################
# Data Helpers
##############################################################################

def _value(row, field):
    """
    Safely retrieve a value from a seed record.
    """

    value = row.get(field)

    if value is None:
        return ""

    return str(value).strip()


def _normalize(text):
    """
    Collapse multiple whitespace characters into single spaces.
    """

    if not text:
        return ""

    return " ".join(str(text).split())


##############################################################################
# Label List Builder
##############################################################################

def _build_label_list(
    label_data,
    include_background=False,
):
    """
    Expand (row, quantity) pairs into a flat list of labels.

    Returns

        [
            (row, False),
            (row, False),
            (row, True),
            ...
        ]
    """

    labels = []

    for row, qty in label_data:

        qty = int(qty)

        ##################################################################
        # Standard labels
        ##################################################################

        for _ in range(qty):
            labels.append((row, False))

        ##################################################################
        # Background labels
        ##################################################################

        if include_background:

            background = _value(
                row,
                "BackgroundInfo",
            )

            if background:

                for _ in range(qty):
                    labels.append((row, True))

    return labels


##############################################################################
# Avery Page Builder
##############################################################################

def _create_page_table(document):
    """
    Creates one Avery 94207 page.

        Label   Gutter   Label

        4.000   0.125    4.000

    Five rows.
    """

    table = document.add_table(
        rows=ROWS,
        cols=3,
    )

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    _remove_table_borders(table)

    ######################################################################
    # Fixed column widths
    ######################################################################

    widths = (
        LABEL_WIDTH,
        GUTTER_WIDTH,
        LABEL_WIDTH,
    )

    for c, width in enumerate(widths):

        table.columns[c].width = Inches(width)

    ######################################################################
    # Configure rows
    ######################################################################

    for row in table.rows:

        _set_row_height(row)

        for c, width in enumerate(widths):

            cell = row.cells[c]

            _clear_cell(cell)

            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.TOP
            )

            _set_cell_width(
                cell,
                width,
            )

            _set_cell_margins(
                cell,
                top=0,
                bottom=0,
                left=0,
                right=0,
            )

    return table


##############################################################################
# Label Cell Lookup
##############################################################################

def _label_cell(
    table,
    label_index,
):
    """
    Returns the proper label cell.

    Label numbers

        0 1
        2 3
        4 5
        6 7
        8 9

    map to

        column 0
        column 2

    because column 1 is the gutter.
    """

    row = label_index // 2

    if label_index % 2 == 0:
        return table.cell(row, 0)

    return table.cell(row, 2)


##############################################################################
# Page Iterator
##############################################################################

def _page_chunks(labels):
    """
    Yield groups of ten labels.
    """

    for start in range(
        0,
        len(labels),
        LABELS_PER_PAGE,
    ):

        yield labels[
            start:start + LABELS_PER_PAGE
        ]


##############################################################################
# Standard Seed Label
##############################################################################

def _draw_seed_label(
    cell,
    row,
):
    """
    Draw one standard CCMGA Seed Library label.

    The label uses ordinary Word paragraphs only.
    No nested tables are created.
    """

    ######################################################################
    # Read seed information
    ######################################################################

    family = _value(row, "Family")
    variety = _value(row, "Variety")

    season = _value(row, "Season")
    edible = _value(row, "Edible")
    year = _value(row, "Year")

    num_seeds = _value(row, "NumSeeds")
    saver = _value(row, "SeedSaverLevel")

    germination = _value(row, "Germination")
    soil_temp = _value(row, "SoilTemperature")

    hybrid = _value(row, "HybridDoNotSave")
    comments = _value(row, "Comments")

    ######################################################################
    # CCMGA heading
    ######################################################################

    _add_text(
        cell,
        "CCMGA Seed Library",
        size=TITLE_SIZE,
        bold=True,
        color=GREEN,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    ######################################################################
    # Family
    ######################################################################

    if family:

        _add_text(
            cell,
            family,
            size=FAMILY_SIZE,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    ######################################################################
    # Variety
    ######################################################################

    if variety:

        _add_text(
            cell,
            variety,
            size=VARIETY_SIZE,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    ######################################################################
    # Hybrid warning
    ######################################################################

    if hybrid:

        _add_text(
            cell,
            hybrid,
            size=SMALL_SIZE,
            bold=True,
            color=RED,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    ######################################################################
    # Comments
    ######################################################################

    if comments:

        _add_text(
            cell,
            _normalize(comments),
            size=SMALL_SIZE,
        )

    ######################################################################
    # Seed information
    ######################################################################

    _add_field(
        cell,
        "Year",
        year,
    )

    _add_field(
        cell,
        "Season",
        season,
    )

    _add_field(
        cell,
        "Edible",
        edible,
    )

    _add_field(
        cell,
        "Seeds",
        num_seeds,
    )

    _add_field(
        cell,
        "Seed Saver",
        saver,
    )

    _add_field(
        cell,
        "Germination",
        germination,
    )

    _add_field(
        cell,
        "Soil Temp",
        soil_temp,
    )



##############################################################################
# Background Information Label
##############################################################################

def _draw_background_label(
    cell,
    row,
):
    """
    Draw a CCMGA Seed Library Background Information label.

    The label uses ordinary paragraphs only.
    No nested tables or divider lines are used.
    """

    ######################################################################
    # Read seed information
    ######################################################################

    family = _value(
        row,
        "Family",
    )

    variety = _value(
        row,
        "Variety",
    )

    background = _value(
        row,
        "BackgroundInfo",
    )

    ######################################################################
    # CCMGA heading
    ######################################################################

    _add_text(
        cell,
        "CCMGA Seed Library",
        size=TITLE_SIZE,
        bold=True,
        color=GREEN,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    ######################################################################
    # Family
    ######################################################################

    if family:

        _add_text(
            cell,
            family,
            size=FAMILY_SIZE,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    ######################################################################
    # Variety
    ######################################################################

    if variety:

        _add_text(
            cell,
            variety,
            size=VARIETY_SIZE,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    ######################################################################
    # Background Information heading
    ######################################################################

    _add_text(
        cell,
        "Background Information",
        size=BODY_SIZE,
        bold=True,
        color=GREEN,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    ######################################################################
    # Background text
    ######################################################################

    if background:

        _add_text(
            cell,
            _normalize(background),
            size=SMALL_SIZE,
            align=WD_ALIGN_PARAGRAPH.LEFT,
        )


##############################################################################
# Document Builder
##############################################################################

def _new_page(document):
    """
    Start a new Word page.

    A page break is added after each completed Avery page,
    except after the final page.
    """

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0

    from docx.enum.text import WD_BREAK
    paragraph.add_run().add_break(WD_BREAK.PAGE)

def _build_document(
    document,
    labels,
):
    """
    Build all Avery 94207 label pages.

    Each page contains:

        5 rows
        2 labels per row
        10 labels per page

    The outer table contains three columns:

        4.000-inch left label
        0.125-inch gutter
        4.000-inch right label
    """

    if not labels:
        return document

    first_page = True

    for page in _page_chunks(labels):

        ##################################################################
        # Start each additional label sheet on a new page
        ##################################################################

        if not first_page:

            _new_page(document)

        first_page = False

        ##################################################################
        # Create one Avery 94207 page
        ##################################################################

        table = _create_page_table(document)

        ##################################################################
        # Draw the labels on this page
        ##################################################################

        for index, (row, is_background) in enumerate(page):

            cell = _label_cell(
                table,
                index,
            )

            if is_background:

                _draw_background_label(
                    cell,
                    row,
                )

            else:

                _draw_seed_label(
                    cell,
                    row,
                )

        ##################################################################
        # Clear unused label positions on the final page
        ##################################################################

        for index in range(
            len(page),
            LABELS_PER_PAGE,
        ):

            cell = _label_cell(
                table,
                index,
            )

            _clear_cell(cell)

            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.TOP
            )

            _set_cell_margins(
                cell,
                top=0,
                bottom=0,
                left=0,
                right=0,
            )

    return document


##############################################################################
# Public DOCX Generator
##############################################################################

def generate_labels_docx(
    label_data,
    include_background=False,
):
    """
    Generate Avery 94207 labels as a Microsoft Word DOCX file.

    Parameters
    ----------
    label_data : list

        A list containing:

            (seed_row, quantity)

        Example:

            [
                (row_1, 2),
                (row_2, 1),
            ]

    include_background : bool

        When True, a Background Information label is added
        for each seed label when BackgroundInfo contains text.

    Returns
    -------
    bytes | None

        Returns the completed DOCX file as bytes.

        Returns None when no labels are selected or when
        the document cannot be generated.
    """

    try:

        ##################################################################
        # Expand seed quantities into individual labels
        ##################################################################

        labels = _build_label_list(
            label_data,
            include_background=include_background,
        )

        if not labels:

            return None

        ##################################################################
        # Create and populate the Word document
        ##################################################################

        document = _setup_document()

        _build_document(
            document,
            labels,
        )

        ##################################################################
        # Save the document to memory
        ##################################################################

        buffer = BytesIO()

        document.save(buffer)

        buffer.seek(0)

        return buffer.getvalue()

    except Exception:

        return None

